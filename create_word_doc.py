"""
Convert the thesis project plan into a styled Word document.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── Colours ──────────────────────────────────────────────────────────────────
DARK_NAVY  = RGBColor(0x1A, 0x1A, 0x2E)
TEAL       = RGBColor(0x00, 0xA8, 0x88)
PURPLE     = RGBColor(0x6C, 0x2A, 0xDD)
AMBER      = RGBColor(0xD4, 0x8A, 0x00)
RED        = RGBColor(0xCC, 0x44, 0x44)
BLUE       = RGBColor(0x00, 0x77, 0xCC)
DARK_GREY  = RGBColor(0x33, 0x33, 0x33)
MID_GREY   = RGBColor(0x66, 0x66, 0x66)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

PHASE_COLOURS = {
    "1": TEAL, "2": TEAL,
    "3": PURPLE, "4": PURPLE,
    "5": AMBER, "6": AMBER,
    "7": RED, "8": RED,
}


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a cell. Each param is (size_eighths, color_hex, style)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            sz, color, style = val
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), style)
            el.set(qn('w:sz'), str(sz))
            el.set(qn('w:color'), color)
            el.set(qn('w:space'), '0')
            borders.append(el)
    tcPr.append(borders)


def add_horizontal_rule(doc, color='00D4AA', thickness=4):
    """Add a coloured horizontal rule."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:color'), color)
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)
    return p


def style_paragraph(p, font_size=11, color=DARK_GREY, bold=False, italic=False,
                    space_before=0, space_after=6, font_name='Calibri', alignment=None):
    """Apply styling to a paragraph."""
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if alignment:
        p.alignment = alignment
    for run in p.runs:
        run.font.size = Pt(font_size)
        run.font.color.rgb = color
        run.font.bold = bold
        run.font.italic = italic
        run.font.name = font_name


# ── Create Document ──────────────────────────────────────────────────────────
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK_GREY

# Set narrow margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════

# Spacing before title
for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('HUMAN vs. AI')
run.font.size = Pt(36)
run.font.color.rgb = DARK_NAVY
run.font.bold = True
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(4)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Analysing the Advancement of AI in Music Mastering')
run.font.size = Pt(18)
run.font.color.rgb = TEAL
run.font.bold = False
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(24)

# Horizontal rule
add_horizontal_rule(doc, '00A888', 6)

# Meta info
meta_lines = [
    ('Student:', 'Michael Adams'),
    ('Programme:', 'MSc Music and Media Technologies'),
    ('Institution:', 'Trinity College Dublin'),
    ('Duration:', '8 Weeks (2 Months)'),
    ('Document:', 'Project Plan'),
]

for label, value in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + '  ')
    run.font.size = Pt(11)
    run.font.color.rgb = MID_GREY
    run.font.bold = True
    run.font.name = 'Calibri'
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GREY
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)

# Page break
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
run = p.add_run('PROJECT OVERVIEW')
run.font.size = Pt(20)
run.font.color.rgb = DARK_NAVY
run.font.bold = True
p.paragraph_format.space_after = Pt(4)

add_horizontal_rule(doc, '00A888', 4)

p = doc.add_paragraph()
run = p.add_run(
    'This project investigates the current state of AI-driven music mastering by conducting '
    'a rigorous blind comparison between professionally mastered tracks and AI-mastered equivalents. '
    'Through structured listening tests with diverse participant groups, the study aims to determine '
    'whether modern AI mastering tools have reached a level of quality that is perceptually comparable '
    'to human mastering engineers.'
)
run.font.size = Pt(11)
run.font.color.rgb = DARK_GREY
p.paragraph_format.space_after = Pt(12)

# Summary Timeline Table
p = doc.add_paragraph()
run = p.add_run('Summary Timeline')
run.font.size = Pt(14)
run.font.color.rgb = DARK_NAVY
run.font.bold = True
p.paragraph_format.space_after = Pt(6)

timeline_data = [
    ('Week', 'Phase', 'Key Activities'),
    ('1', 'Preparation', 'Literature review, song selection, ethics'),
    ('2', 'Preparation', 'Engage engineers, begin AI research'),
    ('3', 'Production', 'Process AI masters, design blind test'),
    ('4', 'Production', 'Normalise audio, build test platform, pilot'),
    ('5', 'Testing', 'Recruit participants, begin blind tests'),
    ('6', 'Testing', 'Complete tests, preliminary analysis'),
    ('7', 'Analysis', 'Deep analysis, begin thesis writing'),
    ('8', 'Analysis', 'Conclusions, final write-up, presentation'),
]

table = doc.add_table(rows=len(timeline_data), cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (wk, phase, activities) in enumerate(timeline_data):
    row = table.rows[i]
    row.cells[0].text = wk
    row.cells[1].text = phase
    row.cells[2].text = activities

    for j, cell in enumerate(row.cells):
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

        if i == 0:
            # Header row
            set_cell_shading(cell, '1A1A2E')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(10)
        else:
            color = PHASE_COLOURS.get(wk, TEAL)
            if i % 2 == 0:
                set_cell_shading(cell, 'F7F7F7')
            if j == 0:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = color

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(2)
    row.cells[1].width = Cm(3)
    row.cells[2].width = Cm(11)

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════════════════════
# WEEK-BY-WEEK DETAIL
# ═══════════════════════════════════════════════════════════════════════════

weeks = [
    {
        'num': '1',
        'title': 'Research, Planning & Song Selection',
        'phase': 'Preparation',
        'objectives': [
            'Finalise the research methodology and evaluation framework',
            'Select and curate a diverse set of pre-mastered songs for the study',
        ],
        'tasks': [
            ('Literature Review & Background Research', [
                'Survey existing academic literature on AI mastering (LANDR, CloudBounce, iZotope Ozone AI, Dolby.io, BandLab, eMastered, etc.)',
                'Review prior comparative studies on human vs. AI audio processing',
                'Document the current landscape of AI mastering technology',
            ]),
            ('Song Selection & Curation', [
                'Select 4\u20136 pre-mastered songs spanning different genres (e.g., rock, electronic, jazz, classical, pop, singer-songwriter)',
                'Ensure songs vary in dynamic range, instrumentation density, and tonal complexity',
                'Source songs where you have permission to use them (original compositions, Creative Commons, or artist-permitted tracks)',
                'Ensure all tracks are high-quality unmastered mixes (minimum 24-bit WAV)',
            ]),
            ('Ethics & Permissions', [
                'Confirm ethical approval requirements with TCD for human listening tests',
                'Prepare participant consent forms and information sheets',
                'Secure written permission for use of any third-party music',
            ]),
        ],
        'deliverables': [
            'Finalised song list with rationale for each selection',
            'Literature review summary document',
            'Ethics/consent documentation drafted',
        ],
    },
    {
        'num': '2',
        'title': 'Engage Professional Mastering Engineers',
        'phase': 'Preparation',
        'objectives': [
            'Identify and engage 2\u20133 professional mastering engineers',
            'Provide them with the selected tracks and clear briefs',
        ],
        'tasks': [
            ('Engineer Identification & Outreach', [
                'Research and shortlist mastering engineers (ideally with varying experience levels and studio setups)',
                'Contact engineers with a clear project brief explaining the academic purpose',
                'Agree on timelines, fees (if applicable), and deliverable formats',
            ]),
            ('Brief Preparation & Delivery', [
                'Prepare a standardised mastering brief for each song (target loudness, format requirements, genre-specific notes)',
                'Provide all engineers with identical source files (same format, sample rate, bit depth)',
                'Request masters in 16-bit/44.1kHz WAV (CD quality) as well as 24-bit/48kHz for analysis',
                'Ask engineers NOT to add any unique identifiers (fades, tags, etc.)',
            ]),
            ('Begin AI Mastering Research', [
                'Identify 2\u20133 AI mastering platforms to test (e.g., LANDR, CloudBounce, iZotope Ozone AI Assistant)',
                'Create accounts and familiarise yourself with each platform\u2019s workflow',
                'Document the feature sets, pricing tiers, and configuration options of each tool',
            ]),
        ],
        'deliverables': [
            'Confirmed list of participating mastering engineers',
            'Mastering briefs sent to all engineers',
            'AI mastering platform shortlist with feature comparison',
        ],
    },
    {
        'num': '3',
        'title': 'AI Mastering & Preliminary Collection',
        'phase': 'Production',
        'objectives': [
            'Master all songs through the selected AI platforms',
            'Begin collecting completed human masters',
            'Design the blind test framework',
        ],
        'tasks': [
            ('AI Mastering Execution', [
                'Process each song through each selected AI mastering platform',
                'Experiment with available settings (genre presets, loudness targets, tonal adjustments)',
                'Select the best AI master from each platform for each song (or use default settings for consistency \u2014 document your decision)',
                'Export all AI masters in matching formats (16-bit/44.1kHz and 24-bit/48kHz WAV)',
            ]),
            ('Documentation of AI Process', [
                'Screenshot and document all settings used for each AI master',
                'Record processing times and any limitations encountered',
                'Note any songs where AI tools struggled or produced unexpected results',
            ]),
            ('Blind Test Design', [
                'Design the listening test methodology (ABX, MUSHRA, or custom scoring approach)',
                'Define scoring criteria (e.g., clarity, punch, warmth, stereo image, dynamics, overall preference)',
                'Use a Likert scale (1\u201310) for each criterion',
                'Plan the listening environment requirements (headphones vs. speakers, room acoustics)',
            ]),
        ],
        'deliverables': [
            'Complete set of AI-mastered tracks (all songs \u00d7 all platforms)',
            'AI mastering process documentation',
            'Blind test methodology document',
        ],
    },
    {
        'num': '4',
        'title': 'Collection, Normalisation & Test Preparation',
        'phase': 'Production',
        'objectives': [
            'Collect all human masters',
            'Normalise and prepare all tracks for blind testing',
            'Build the listening test infrastructure',
        ],
        'tasks': [
            ('Collect Human Masters', [
                'Follow up with engineers and collect all completed masters',
                'Verify file formats and quality match specifications',
                'Address any issues or request revisions if needed',
            ]),
            ('Audio Normalisation & Preparation', [
                'Loudness-normalise all masters (human and AI) to a consistent level (e.g., \u221214 LUFS integrated)',
                'This is critical \u2014 loudness differences bias listeners toward louder versions',
                'Trim all tracks to identical start/end points',
                'Randomise and assign anonymous labels (e.g., \u201cVersion A\u201d, \u201cVersion B\u201d, \u201cVersion C\u201d)',
                'Create a master reference document mapping labels to sources (keep strictly confidential)',
            ]),
            ('Test Platform Setup & Pilot', [
                'Set up the listening test platform (webMUSHRA, Google Forms, or custom web app)',
                'Upload all anonymised audio files',
                'Create the questionnaire with scoring criteria and demographic questions',
                'Run a pilot test with 2\u20133 people to validate the flow',
                'Refine based on pilot feedback',
            ]),
        ],
        'deliverables': [
            'All mastered tracks collected, normalised, and anonymised',
            'Functional listening test platform',
            'Pilot test completed with feedback incorporated',
        ],
    },
    {
        'num': '5',
        'title': 'Recruit Participants & Begin Blind Testing',
        'phase': 'Testing',
        'objectives': [
            'Recruit a diverse pool of listeners',
            'Begin conducting blind listening tests',
        ],
        'tasks': [
            ('Participant Recruitment', [
                'Recruit from MMT classmates (trained ears, music technology background)',
                'Recruit from other TCD music students (performance, composition, musicology)',
                'Engage professional musicians/engineers (external contacts)',
                'Include general listeners (non-musicians for broader perspective)',
                'Aim for a minimum of 20\u201330 participants across groups',
                'Distribute consent forms and collect signed agreements',
            ]),
            ('Listening Test Sessions', [
                'Schedule and conduct listening test sessions',
                'Provide consistent listening conditions where possible',
                'For remote participants, provide clear instructions on playback requirements',
                'Monitor for test fatigue \u2014 keep individual sessions under 30\u201340 minutes',
            ]),
            ('Data Collection', [
                'Record all responses in a structured format (spreadsheet/database)',
                'Track participant metadata (group, experience level, listening setup)',
                'Note any technical issues or anomalies during sessions',
            ]),
        ],
        'deliverables': [
            'Minimum 15+ completed listening tests by end of week',
            'Organised raw data spreadsheet',
            'Participant log with demographics',
        ],
    },
    {
        'num': '6',
        'title': 'Complete Testing & Begin Analysis',
        'phase': 'Testing',
        'objectives': [
            'Complete all remaining listening tests',
            'Begin preliminary data analysis',
        ],
        'tasks': [
            ('Complete Remaining Tests', [
                'Follow up with outstanding participants',
                'Conduct any remaining in-person sessions',
                'Close the test window and finalise data collection',
            ]),
            ('Data Cleaning & Organisation', [
                'Clean and validate all collected data',
                'Remove incomplete responses or flag outliers',
                'Organise data by participant group, song, and mastering source',
            ]),
            ('Preliminary Statistical Analysis', [
                'Calculate mean scores and standard deviations for each version of each song',
                'Compare human vs. AI scores across all evaluation criteria',
                'Break down results by participant group (expert vs. general listener)',
                'Generate initial charts and visualisations',
                'Consider statistical tests: paired t-tests, ANOVA, or Wilcoxon signed-rank tests',
            ]),
        ],
        'deliverables': [
            'Complete dataset (all participants)',
            'Preliminary analysis with initial charts',
            'Identified trends and patterns',
        ],
    },
    {
        'num': '7',
        'title': 'Deep Analysis & Thesis Writing',
        'phase': 'Analysis',
        'objectives': [
            'Conduct comprehensive statistical analysis',
            'Begin writing the thesis/report',
        ],
        'tasks': [
            ('Comprehensive Analysis', [
                'Overall preference: Human vs. AI',
                'Per-criterion analysis (clarity, dynamics, warmth, etc.)',
                'Genre-dependent performance (does AI do better with certain genres?)',
                'Expert vs. non-expert perception differences',
                'Platform-specific performance (which AI tool performed best?)',
                'Calculate statistical significance for key findings',
                'Create publication-quality charts, graphs, and tables',
            ]),
            ('Spectral & Technical Analysis (Optional Enhancement)', [
                'Perform spectral analysis comparing human and AI masters (frequency response, dynamic range, stereo width)',
                'Use tools like iZotope Insight, SPAN, or Python (librosa) for technical measurements',
                'Correlate technical differences with perceptual scores',
            ]),
            ('Begin Thesis Writing', [
                'Introduction and literature review',
                'Methodology chapter (detailed description of the entire process)',
                'Results chapter (present findings with supporting data)',
            ]),
        ],
        'deliverables': [
            'Complete statistical analysis with all charts and tables',
            'Draft of Introduction, Literature Review, and Methodology chapters',
            'Technical analysis report (if included)',
        ],
    },
    {
        'num': '8',
        'title': 'Conclusions, Final Write-Up & Presentation',
        'phase': 'Completion',
        'objectives': [
            'Draw conclusions from the analysis',
            'Complete the thesis document',
            'Prepare the final presentation',
        ],
        'tasks': [
            ('Conclusions & Discussion', [
                'Has AI mastering reached professional quality?',
                'Where does AI excel and where does it fall short?',
                'What are the implications for the music industry?',
                'What are the limitations of this study?',
                'Discuss future directions and recommendations',
            ]),
            ('Complete Thesis Document', [
                'Write the Discussion and Conclusions chapters',
                'Write the Abstract',
                'Compile references and bibliography',
                'Proofread and format the entire document',
                'Ensure all appendices are included (consent forms, raw data, test materials)',
            ]),
            ('Final Presentation & Submission', [
                'Create a presentation summarising the project and key findings',
                'Rehearse the presentation',
                'Prepare for Q&A from supervisors and lecturers',
                'Final proofread, quality check, and submit',
            ]),
        ],
        'deliverables': [
            'Completed thesis document',
            'Final presentation',
            'All supporting materials and raw data archived',
        ],
    },
]

for week in weeks:
    doc.add_page_break()

    wk_num = week['num']
    color = PHASE_COLOURS.get(wk_num, TEAL)

    # Week header
    p = doc.add_paragraph()
    run = p.add_run(f'WEEK {wk_num}')
    run.font.size = Pt(12)
    run.font.color.rgb = color
    run.font.bold = True
    run = p.add_run(f'   |   {week["phase"].upper()}')
    run.font.size = Pt(12)
    run.font.color.rgb = MID_GREY
    run.font.bold = True
    p.paragraph_format.space_after = Pt(2)

    # Week title
    p = doc.add_paragraph()
    run = p.add_run(week['title'])
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_NAVY
    run.font.bold = True
    p.paragraph_format.space_after = Pt(4)

    # Coloured rule
    color_hex = f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
    add_horizontal_rule(doc, color_hex, 4)

    # Objectives
    p = doc.add_paragraph()
    run = p.add_run('Objectives')
    run.font.size = Pt(13)
    run.font.color.rgb = color
    run.font.bold = True
    p.paragraph_format.space_after = Pt(4)

    for obj in week['objectives']:
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        run = p.add_run(obj)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GREY
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()  # spacer

    # Tasks
    p = doc.add_paragraph()
    run = p.add_run('Tasks')
    run.font.size = Pt(13)
    run.font.color.rgb = color
    run.font.bold = True
    p.paragraph_format.space_after = Pt(4)

    for task_num, (task_title, subtasks) in enumerate(week['tasks'], 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{task_num}. {task_title}')
        run.font.size = Pt(12)
        run.font.color.rgb = DARK_NAVY
        run.font.bold = True
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)

        for subtask in subtasks:
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            run = p.add_run(subtask)
            run.font.size = Pt(10.5)
            run.font.color.rgb = DARK_GREY
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(1.5)

    doc.add_paragraph()  # spacer

    # Deliverables
    p = doc.add_paragraph()
    run = p.add_run('Deliverables')
    run.font.size = Pt(13)
    run.font.color.rgb = color
    run.font.bold = True
    p.paragraph_format.space_after = Pt(4)

    for d in week['deliverables']:
        p = doc.add_paragraph()
        run = p.add_run(f'\u2713  {d}')
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GREY
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)

# ═══════════════════════════════════════════════════════════════════════════
# RISK MITIGATION
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()

p = doc.add_paragraph()
run = p.add_run('RISK MITIGATION')
run.font.size = Pt(20)
run.font.color.rgb = DARK_NAVY
run.font.bold = True
p.paragraph_format.space_after = Pt(4)

add_horizontal_rule(doc, 'CC4444', 4)

risks = [
    ('Risk', 'Impact', 'Mitigation'),
    ('Engineer Delays', 'Delays Week 4 tasks', 'Engage engineers early; have backup engineers identified'),
    ('Low Participant Count', 'Weakens statistical significance', 'Start recruitment in Week 3; leverage social media and course networks'),
    ('AI Platform Limitations', 'Incomplete AI dataset', 'Test platforms early in Week 2; have backup platforms ready'),
    ('Test Platform Issues', 'Delays testing', 'Pilot test in Week 4; have offline backup (USB + paper forms)'),
    ('Ethics Approval Delays', 'Blocks testing phase', 'Submit ethics application in Week 1; follow up proactively'),
]

table = doc.add_table(rows=len(risks), cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (risk, impact, mitigation) in enumerate(risks):
    row = table.rows[i]
    row.cells[0].text = risk
    row.cells[1].text = impact
    row.cells[2].text = mitigation

    for j, cell in enumerate(row.cells):
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

        if i == 0:
            set_cell_shading(cell, '1A1A2E')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
        else:
            if i % 2 == 0:
                set_cell_shading(cell, 'F7F7F7')
            if j == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RED

for row in table.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(4)
    row.cells[2].width = Cm(8.5)

# ═══════════════════════════════════════════════════════════════════════════
# TOOLS & RESOURCES
# ═══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('TOOLS & RESOURCES')
run.font.size = Pt(20)
run.font.color.rgb = DARK_NAVY
run.font.bold = True
p.paragraph_format.space_after = Pt(4)

add_horizontal_rule(doc, '00A888', 4)

tool_sections = [
    ('AI Mastering Platforms', [
        'LANDR (landr.com)',
        'CloudBounce (cloudbounce.com)',
        'eMastered (emastered.com)',
        'iZotope Ozone AI Assistant',
        'BandLab Mastering (bandlab.com)',
        'Dolby.io Media Mastering',
    ]),
    ('Analysis Tools', [
        'Python (librosa, scipy, matplotlib) for audio and statistical analysis',
        'SPSS or R for advanced statistics',
        'iZotope Insight / Voxengo SPAN for spectral analysis',
        'Excel / Google Sheets for data organisation',
    ]),
    ('Listening Test Platforms', [
        'webMUSHRA (open-source, designed for audio quality evaluation)',
        'Google Forms + cloud-hosted audio',
        'Custom web application (if needed)',
    ]),
]

for section_title, items in tool_sections:
    p = doc.add_paragraph()
    run = p.add_run(section_title)
    run.font.size = Pt(13)
    run.font.color.rgb = TEAL
    run.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK_GREY
        p.paragraph_format.space_after = Pt(1)

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = "/Users/michaeladams/Work/Esker/Development/AI/Claude/thesis-plan/Thesis_Project_Plan.docx"
doc.save(output_path)
print(f"Word document saved to: {output_path}")
